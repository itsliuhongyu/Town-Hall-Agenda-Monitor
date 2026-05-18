import os
import csv
import json
import re
from openai import OpenAI

# --- File readers ---
def read_pdf(path):
    try:
        import pdfplumber
        text = ""
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text(x_tolerance=1, y_tolerance=3)
                if t:
                    text += t + "\n"
        return text.strip()
    except Exception as e:
        return f"[PDF read error: {e}]"

def read_docx(path):
    try:
        from docx import Document
        doc = Document(path)
        parts = []

        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts)
    except Exception as e:
        return f"[DOCX read error: {e}]"

def read_html(path):
    try:
        from bs4 import BeautifulSoup
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            soup = BeautifulSoup(f.read(), "html.parser")

        for tag in soup(["script", "style"]):
            tag.decompose()

        return soup.get_text(separator="\n").strip()
    except Exception as e:
        return f"[HTML read error: {e}]"

def read_file_content(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return read_pdf(path)
    elif ext == ".docx":
        return read_docx(path)
    elif ext in (".html", ".htm"):
        return read_html(path)
    return None

# --- LM Studio client ---
client = OpenAI(
    base_url="http://localhost:1234/v1",
    api_key="lm-studio"
)

MODEL_NAME = "qwen3.5-9b"

SYSTEM_PROMPT = """
You are extracting agenda items for a newsroom.

Return a JSON object only.

Required schema:
{
  "items": [
    {
      "item": "short agenda item name",
      "importance": "low",
      "original_text": "verbatim or near-verbatim source text"
    }
  ]
}

Extract actual agenda items only.

Do extract:
- numbered items
- lettered items
- consent agenda items
- public hearing items
- old business
- new business
- resolutions
- ordinances
- contracts
- reports
- infrastructure updates
- public facility closures
- public works updates
- action items
- discussion items
- closed session items

Do NOT extract:
- ADA accommodation notices
- open meetings boilerplate
- legal disclaimers
- contact information
- publication notices
- page headers
- page footers
- lines saying no action will be taken
- routine items such as call to order, roll call, pledge, agenda approval, minutes approval, public comment, pay bills, or adjournment

Importance must be exactly one of: low, medium, high.

Use high only for:
- major spending
- rezoning
- development projects
- lawsuits
- tax increases
- emergencies
- major policy changes
- major public infrastructure projects
- public facility closures or disruptions
- major public works projects

Use medium only for:
- contracts
- public hearings
- ordinance changes
- staff reports with possible action
- committee discussions with possible action
- resolutions
- agreements
- purchases
- appointments
- public infrastructure updates
- public works updates
- municipal building system work
- harbor, road, water, sewer, launch, park, or town facility updates

Use low for:
- routine reports
- informational updates without public impact

Do not explain.
Do not include analysis.
Do not include thinking.
"""

SYSTEM_PROMPT_FALLBACK = """
Return JSON only.

Schema:
{
  "items": [
    {
      "item": "short agenda item name",
      "importance": "low",
      "original_text": "verbatim or near-verbatim source text"
    }
  ]
}

Extract actual agenda items only. Do not extract ADA notices, open meetings notices, legal disclaimers, contact information, boilerplate, call to order, roll call, pledge, agenda approval, minutes approval, public comment, pay bills, or adjournment.
"""

ROUTINE_LOW_PATTERNS = [
    r"\bcall (the )?meeting to order\b",
    r"\bcall to order\b",
    r"\broll call\b",
    r"\bpledge of allegiance\b",
    r"\bapprove/?modify agenda\b",
    r"\bapproval of agenda\b",
    r"\bapprove .*minutes\b",
    r"\bapproval of .*minutes\b",
    r"\bpublic comment\b",
    r"\bpublic concerns? ?&? comments?\b",
    r"\bitems to share\b",
    r"\bclerk'?s? reports?\b",
    r"\btreasurer'?s? reports?\b",
    r"\bapproval to pay bills\b",
    r"\bpay bills\b",
    r"\badjourn(ment)?\b",
]

def sanitize_text(text):
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E]", " ", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

def normalize_line(line):
    return re.sub(r"\s+", " ", line).strip()

def is_routine_low_value(text):
    lower = normalize_line(text).lower()
    return any(
        re.search(pattern, lower, re.IGNORECASE)
        for pattern in ROUTINE_LOW_PATTERNS
    )

def is_boilerplate(text):
    lower = normalize_line(text).lower()

    boilerplate_patterns = [
        r"no action will be taken",
        r"governmental body specifically referred to",
        r"above stated meeting",
        r"americans with disabilities act",
        r"\bada\b",
        r"alternative format",
        r"larger print",
        r"audio tapes?",
        r"braille",
        r"reader interpreters?",
        r"amplifiers?",
        r"transcription",
        r"48[- ]?hours? prior",
        r"48[- ]?hours? after request",
        r"delivery of that alternative format",
        r"notify the town",
        r"contact .* at \(?\d{3}\)?[- ]?\d{3}[- ]?\d{4}",
        r"notice is hereby given",
        r"posted this",
        r"published this",
        r"agenda is subject to change",
        r"members of other governmental bodies",
        r"may be present",
        r"for information only",
        r"accommodations?",
    ]

    return any(re.search(pattern, lower, re.IGNORECASE) for pattern in boilerplate_patterns)

def extract_message_text(message):
    content = getattr(message, "content", None)
    reasoning = getattr(message, "reasoning_content", None)

    if content and content.strip():
        return content.strip()

    if reasoning and reasoning.strip():
        return reasoning.strip()

    return ""

def classify_importance(text):
    lower = text.lower()

    if is_boilerplate(lower) or is_routine_low_value(lower):
        return "low"

    high_terms = [
        "rezoning", "rezone", "zoning map", "conditional use",
        "development agreement", "development project", "subdivision",
        "tax levy", "tax increase", "bond", "borrowing",
        "lawsuit", "litigation", "emergency", "major policy",
        "comprehensive plan amendment", "tif", "tax incremental",
        "capital project",
        "facility closure", "temporarily closed", "closed for repairs",
        "road closure", "bridge closure",
        "major infrastructure", "major public works",
    ]

    medium_terms = [
        "contract", "agreement", "public hearing", "ordinance",
        "resolution", "staff report", "possible action", "discussion and action",
        "may act", "committee", "purchase", "proposal", "bid",
        "appointment", "license", "permit", "memorandum of understanding",
        "mou", "lease", "grant",
        "dredging", "harbor", "marina", "boat launch", "launch",
        "geothermal", "installation", "repairs", "construction",
        "public works", "infrastructure", "town hall", "municipal building",
        "water", "sewer", "stormwater", "road", "street", "bridge",
        "park", "dock", "wells", "piping", "channel", "soundings",
        "public facility", "maintenance project"
    ]

    low_terms = [
        "announcements", "correspondence",
        "informational", "schedule"
    ]

    if any(term in lower for term in high_terms):
        return "high"

    if any(term in lower for term in medium_terms):
        return "medium"

    if any(term in lower for term in low_terms):
        return "low"

    if "report" in lower:
        return "low"

    return "low"

def short_item_name(original_text):
    text = normalize_line(original_text)

    text = re.sub(r"^(item\s*)?[\dIVXivx]+[\.\)]\s*", "", text)
    text = re.sub(r"^[A-Z][\.\)]\s*", "", text)
    text = re.sub(r"^[a-z][\.\)]\s*", "", text)
    text = re.sub(r"^\(?\d+\)?\s*", "", text)
    text = text.strip(" -:;")

    if len(text) > 140:
        text = text[:137].rstrip() + "..."

    return text or normalize_line(original_text)

def split_into_chunks(text, max_chars=5500, overlap=500):
    if len(text) <= max_chars:
        return [text]

    chunks = []
    start = 0

    while start < len(text):
        end = min(start + max_chars, len(text))

        if end < len(text):
            newline = text.rfind("\n", start, end)
            if newline > start + max_chars // 2:
                end = newline

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        if end >= len(text):
            break

        start = max(0, end - overlap)

    return chunks

def filter_text_before_llm(text):
    lines = [line for line in text.splitlines()]
    kept = []

    for line in lines:
        clean = normalize_line(line)

        if not clean:
            kept.append(line)
            continue

        if is_boilerplate(clean):
            continue

        if is_routine_low_value(clean):
            continue

        kept.append(line)

    return "\n".join(kept).strip()

def call_llm(filename, content, system_prompt=None):
    if system_prompt is None:
        system_prompt = SYSTEM_PROMPT

    content = filter_text_before_llm(content)

    if not content:
        return '{"items": []}'

    user_msg = f"""Filename: {filename}

Agenda content:
{content}

Return JSON only. Start your answer with {{ and end with }}."""

    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_msg}
        ],
        temperature=0.1,
        top_p=0.9,
        max_tokens=1800,
    )

    msg = response.choices[0].message
    raw = extract_message_text(msg)

    if not raw:
        print("    [Debug] Empty LM Studio response object:")
        print(msg)
        return ""

    return raw

def parse_llm_response(raw):
    cleaned = raw.strip()

    cleaned = re.sub(r"<think>.*?</think>", "", cleaned, flags=re.DOTALL).strip()
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned).strip()

    match = re.search(r"\{.*\}", cleaned, flags=re.DOTALL)
    if not match:
        raise json.JSONDecodeError("No JSON object found", cleaned, 0)

    data = json.loads(match.group(0))
    items = data.get("items", [])
    parsed = []

    for item in items:
        item_name = str(item.get("item", "")).strip()
        original_text = str(item.get("original_text", "")).strip()
        importance = str(item.get("importance", "low")).strip().lower()

        if not item_name and original_text:
            item_name = short_item_name(original_text)

        if not original_text and item_name:
            original_text = item_name

        if not item_name:
            continue

        combined = item_name + " " + original_text

        if is_boilerplate(combined):
            continue

        if is_routine_low_value(combined):
            continue

        if importance not in ("low", "medium", "high"):
            importance = classify_importance(combined)
        else:
            rule_importance = classify_importance(combined)

            if rule_importance == "high":
                importance = "high"
            elif rule_importance == "medium" and importance == "low":
                importance = "medium"

        parsed.append({
            "item": item_name,
            "importance": importance,
            "original_text": original_text
        })

    return parsed

def extract_agenda_items_by_rules(content):
    lines = [normalize_line(line) for line in content.splitlines()]
    lines = [line for line in lines if line]

    items = []
    current = None

    agenda_marker = re.compile(
        r"""^(
            (\d+|[A-Z]|[a-z]|[IVXivx]+)[\.\)]\s+
            |
            \(?\d+\)?\s+
            |
            item\s+\d+[:\.\)]?\s+
        )""",
        re.VERBOSE
    )

    actionish = re.compile(
        r"\b(approve|approval|consider|discussion|possible action|public hearing|"
        r"resolution|ordinance|contract|agreement|permit|license|appointment|"
        r"rezoning|zoning|bid|purchase|grant|closed session|"
        r"new business|old business|consent agenda|claims|vouchers|"
        r"dredging|harbor|marina|boat launch|geothermal|installation|repairs|"
        r"construction|infrastructure|public works|town hall|municipal building|"
        r"road|street|bridge|water|sewer|stormwater|dock|wells|piping|channel)\b",
        re.IGNORECASE
    )

    weak_report_only = re.compile(
        r"^\s*(staff|committee|department|administrator|chair|treasurer|clerk)?\s*reports?\s*$",
        re.IGNORECASE
    )

    for line in lines:
        if len(line) < 3:
            continue

        if is_boilerplate(line) or is_routine_low_value(line):
            if current:
                items.append(current.strip())
                current = None
            continue

        starts_like_item = bool(agenda_marker.match(line))
        contains_action = bool(actionish.search(line))

        if weak_report_only.match(line):
            contains_action = False

        if starts_like_item or contains_action:
            if current:
                items.append(current.strip())

            current = line
        else:
            if current and len(line) < 250:
                if not is_boilerplate(line) and not is_routine_low_value(line):
                    current += " " + line

    if current:
        items.append(current.strip())

    cleaned_items = []

    skip_patterns = [
        r"^page \d+",
        r"^\d+/\d+/\d+",
        r"^https?://",
        r"^www\.",
        r"^agenda packet$",
        r"^notice$",
    ]

    for original in items:
        original = normalize_line(original)

        if is_boilerplate(original) or is_routine_low_value(original):
            continue

        if any(re.search(p, original, re.IGNORECASE) for p in skip_patterns):
            continue

        if len(original) < 5:
            continue

        cleaned_items.append({
            "item": short_item_name(original),
            "importance": classify_importance(original),
            "original_text": original
        })

    return cleaned_items

def dedupe_items(items):
    seen = set()
    deduped = []

    for item in items:
        original_text = normalize_line(item.get("original_text", ""))
        item_name = normalize_line(item.get("item", ""))

        combined = item_name + " " + original_text

        if is_boilerplate(combined) or is_routine_low_value(combined):
            continue

        key = re.sub(r"[^a-z0-9]+", " ", (original_text or item_name).lower()).strip()

        if not key or key in seen:
            continue

        seen.add(key)

        importance = item.get("importance", "low")
        rule_importance = classify_importance(combined)

        if rule_importance == "high":
            importance = "high"
        elif rule_importance == "medium" and importance == "low":
            importance = "medium"

        if importance not in ("low", "medium", "high"):
            importance = rule_importance

        deduped.append({
            "item": item_name or short_item_name(original_text),
            "importance": importance,
            "original_text": original_text or item_name
        })

    return deduped

def analyze_agenda(filename, content):
    content = sanitize_text(content)

    all_items = []

    rule_items = extract_agenda_items_by_rules(content)
    all_items.extend(rule_items)

    llm_content = filter_text_before_llm(content)

    if llm_content:
        chunks = split_into_chunks(llm_content, max_chars=5500, overlap=500)

        for idx, chunk in enumerate(chunks, start=1):
            raw = ""

            try:
                print(f"    LLM chunk {idx}/{len(chunks)}...")
                raw = call_llm(filename, chunk, system_prompt=SYSTEM_PROMPT)
                llm_items = parse_llm_response(raw)
                all_items.extend(llm_items)

            except json.JSONDecodeError as e:
                print(f"    [Bad JSON on chunk {idx}: {e}]")

                if raw:
                    print(f"    [Returned: {raw[:300]}]")

                try:
                    raw = call_llm(filename, chunk[:2500], system_prompt=SYSTEM_PROMPT_FALLBACK)
                    llm_items = parse_llm_response(raw)
                    all_items.extend(llm_items)
                except Exception as e2:
                    print(f"    [Fallback failed on chunk {idx}: {e2}]")

            except Exception as e:
                print(f"    [LLM error on chunk {idx}: {e}]")

    all_items = dedupe_items(all_items)

    if all_items:
        return all_items

    return [{
        "item": "[No agenda items extracted]",
        "importance": "low",
        "original_text": content[:500]
    }]

# --- Main ---
def main():
    temps_dir = "./temps"
    output_dir = "."

    supported_exts = {".pdf", ".docx", ".html", ".htm"}

    if not os.path.isdir(temps_dir):
        print(f"Missing folder: {temps_dir}")
        return

    # Rename files that have ?v=... query strings in their names
    for f in os.listdir(temps_dir):
        if "?" in f:
            clean_name = f.split("?")[0]
            old_path = os.path.join(temps_dir, f)
            new_path = os.path.join(temps_dir, clean_name)
            if not os.path.exists(new_path):
                os.rename(old_path, new_path)
                print(f"Renamed: {f} -> {clean_name}")

    files = [
        f for f in os.listdir(temps_dir)
        if os.path.splitext(f)[1].lower() in supported_exts
    ]

    if not files:
        print("No supported agenda files found in ./temps")
        return

    print(f"Found {len(files)} random agenda file(s) to process.\n")

    rows_high = []
    rows_medium = []
    rows_low = []

    for filename in sorted(files):
        path = os.path.join(temps_dir, filename)
        print(f"Reading:   {filename}")

        content = read_file_content(path)

        if content is None:
            print("  Skipped unsupported type\n")
            continue

        if not content or content.startswith("["):
            print(f"  Warning: {content}\n")
            rows_low.append({
                "filename": filename,
                "item": content or "[No readable content]",
                "importance": "low",
                "original_text": content or "[No readable content]"
            })
            continue

        print("  Analyzing...")
        items = analyze_agenda(filename, content)

        for item in items:
            if is_boilerplate(item.get("original_text", "")) or is_routine_low_value(item.get("original_text", "")):
                continue

            row = {
                "filename": filename,
                "item": item["item"],
                "importance": item["importance"],
                "original_text": item.get("original_text", "")
            }

            print(f"  - {item['importance'].upper()}: {item['item']}")

            if item["importance"] == "high":
                rows_high.append(row)
            elif item["importance"] == "medium":
                rows_medium.append(row)
            else:
                rows_low.append(row)

        print()

    fieldnames = ["filename", "item", "importance", "original_text"]

    def write_csv(rows, filepath):
        with open(filepath, "w", newline="", encoding="utf-8") as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(rows)
        print(f"Saved {len(rows)} item(s) -> {filepath}")

    write_csv(rows_high, os.path.join(output_dir, "high.csv"))
    write_csv(rows_medium, os.path.join(output_dir, "medium.csv"))
    write_csv(rows_low, os.path.join(output_dir, "low.csv"))

    print(f"\nDone. High: {len(rows_high)}, Medium: {len(rows_medium)}, Low: {len(rows_low)}")

if __name__ == "__main__":
    main()